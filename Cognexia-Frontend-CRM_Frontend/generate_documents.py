"""
Generate PDF and Word documents from COMPLETE_FEATURE_LIST.md with CognexiaAI branding
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document():
    """Create a Word document with CognexiaAI branding"""
    doc = Document()
    
    # Add logo/header section
    header_section = doc.sections[0]
    header = header_section.header
    header_para = header.paragraphs[0]
    header_para.text = "CognexiaAI ERP"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.font.size = Pt(20)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
    
    # Add a subtitle
    subtitle = header.add_paragraph()
    subtitle.text = "Industry 5.0 Enterprise CRM Platform"
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray color
    
    # Read the markdown file
    with open('COMPLETE_FEATURE_LIST.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse markdown and add to document
    lines = content.split('\n')
    
    for line in lines:
        if line.startswith('# '):
            # Heading 1
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        elif line.startswith('## '):
            # Heading 2
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(51, 51, 153)
        elif line.startswith('### '):
            # Heading 3
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.color.rgb = RGBColor(102, 102, 102)
        elif line.startswith('---'):
            # Horizontal line (add empty paragraph as separator)
            doc.add_paragraph()
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        elif line.strip():
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save the document
    doc.save('COMPLETE_FEATURE_LIST.docx')
    print("✓ Word document created: COMPLETE_FEATURE_LIST.docx")

def create_pdf_document():
    """Create a PDF document with CognexiaAI branding using reportlab"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    # Read the markdown file
    with open('COMPLETE_FEATURE_LIST.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create PDF
    pdf_file = 'COMPLETE_FEATURE_LIST.pdf'
    doc = SimpleDocTemplate(pdf_file, pagesize=A4,
                           topMargin=1*inch, bottomMargin=1*inch,
                           leftMargin=1*inch, rightMargin=1*inch)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=HexColor('#0066cc'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        textColor=HexColor('#666666'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=HexColor('#0066cc'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=HexColor('#333399'),
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )
    
    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=HexColor('#666666'),
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=4,
        bulletIndent=10
    )
    
    # Add header
    elements.append(Paragraph('CognexiaAI ERP', title_style))
    elements.append(Paragraph('Industry 5.0 Enterprise CRM Platform', subtitle_style))
    elements.append(Paragraph('"Cognition Meets Precision"', subtitle_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Parse markdown and add to PDF
    lines = content.split('\n')
    
    for line in lines:
        line = line.replace('\r', '')
        if line.startswith('# '):
            # Heading 1
            elements.append(Spacer(1, 0.2*inch))
            elements.append(Paragraph(line[2:], heading1_style))
        elif line.startswith('## '):
            # Heading 2
            elements.append(Paragraph(line[3:], heading2_style))
        elif line.startswith('### '):
            # Heading 3
            elements.append(Paragraph(line[4:], heading3_style))
        elif line.startswith('---'):
            # Horizontal line (add space)
            elements.append(Spacer(1, 0.1*inch))
        elif line.startswith('- '):
            # Bullet point
            bullet_text = '• ' + line[2:]
            elements.append(Paragraph(bullet_text, bullet_style))
        elif line.strip():
            # Regular paragraph
            elements.append(Paragraph(line, styles['Normal']))
    
    # Build PDF
    doc.build(elements)
    print("✓ PDF document created: COMPLETE_FEATURE_LIST.pdf")
    return True

if __name__ == '__main__':
    print("Generating CognexiaAI branded documents...\n")
    
    # Create Word document
    try:
        create_word_document()
    except Exception as e:
        print(f"✗ Error creating Word document: {e}")
    
    # Create PDF document
    try:
        create_pdf_document()
    except Exception as e:
        print(f"✗ Error creating PDF document: {e}")
    
    print("\n✓ Document generation complete!")
